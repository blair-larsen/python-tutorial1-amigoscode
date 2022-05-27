# https://www.youtube.com/watch?v=mJEpimi_tFo

from pydoc import Doc, doc
from docx import Document
from docx.shared import Inches

document = Document()

# profile picture
document.add_picture(
    'me1.jpg', 
    width=Inches(1.0)
)

#  personell info

name = input('What is your name?')
phone = input('Phone Number')
email = input('email')

document.add_paragraph(
    name + '|' + phone + '|' + email)

# about me
document.add_heading('About me')
document.add_paragraph(
    input("Tell about yourself")
)

# work experiance
document.add_heading('Work experiance')
p = document.add_paragraph()

company = input('Enter Company')
start_date = input('start date')
end_date = input('end date')

p.add_run(company + ' ').bold = True
p.add_run(start_date + '-' + end_date + '\n').italic = True

experiance = input('Enter experiance at ' + company)
p.add_run(experiance)

# more experiance
while True:
    more_experiances = input(
        'do you have more ? Yes or No')
    if more_experiances.upper() == 'YES':
        p = document.add_paragraph()

        company = input('Enter Company')
        start_date = input('start date')
        end_date = input('end date')

        p.add_run(company + ' ').bold = True
        p.add_run(start_date + '-' + end_date + '\n').italic = True

        experiance = input('Enter experiance at ' + company)
        p.add_run(experiance)
    else:
        break

# skills
document.add_heading('Skills')
skill = input('enter skills')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    more_skills = input(' do you have more skils? yes or no')
    if more_skills.upper() == 'YES':
        skill = input('enter skills')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break
# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = 'CV generated'

document.save('cv.docx')